"""Production-ready document export service for Bimo.

Provides a canonical Markdown-first parsing and rendering engine that produces:
1. Valid UTF-8 Markdown documents (.md)
2. Professional text-based PDF documents (.pdf) via ReportLab
3. Editable Microsoft Word documents (.docx) via python-docx

Both PDF and DOCX converters consume a unified block/inline representation
derived from the single canonical Markdown source string.
"""

from __future__ import annotations

import io
import re
import unicodedata
import xml.sax.saxutils as saxutils
from dataclasses import dataclass, field
from datetime import datetime, timezone
from typing import List, Optional

# --- Data Classes for Unified Document Representation ---

@dataclass
class InlineSpan:
    text: str
    bold: bool = False
    italic: bool = False
    code: bool = False
    math: bool = False
    strike: bool = False
    link_url: Optional[str] = None


@dataclass
class Block:
    pass


@dataclass
class HeadingBlock(Block):
    level: int
    text: str


@dataclass
class ParagraphBlock(Block):
    text: str


@dataclass
class CodeBlock(Block):
    code: str
    language: str = ""
    is_math: bool = False


@dataclass
class MathBlock(Block):
    latex: str


@dataclass
class BlockquoteBlock(Block):
    lines: List[str] = field(default_factory=list)


@dataclass
class ListItem:
    text: str
    level: int = 0
    number: Optional[int] = None  # None for unordered, int for ordered


@dataclass
class ListBlock(Block):
    ordered: bool
    items: List[ListItem] = field(default_factory=list)


@dataclass
class TableBlock(Block):
    headers: List[str] = field(default_factory=list)
    alignments: List[str] = field(default_factory=list)  # 'left', 'center', 'right'
    rows: List[List[str]] = field(default_factory=list)


@dataclass
class HorizontalRuleBlock(Block):
    pass


# --- Inline Parser & Unicode / Math Normalizers ---

def normalize_unicode_for_export(text: str) -> str:
    """Normalize problematic Unicode characters that fail in standard PDF document fonts.

    Replaces non-breaking hyphens, special dashes, and non-standard symbols with their
    clean, renderable ASCII/typography equivalents so fonts never draw missing-glyph black boxes (■).
    """
    if not text:
        return ""

    replacements = {
        "‑": "-",   # non-breaking hyphen
        "‐": "-",   # hyphen
        "‒": "-",   # figure dash
        "–": "–",   # en dash
        "—": "—",   # em dash
        "―": "—",   # horizontal bar
        "−": "-",   # minus sign
        "­": "",    # soft hyphen
        " ": " ",   # non-breaking space
        "​": "",    # zero-width space
        "‌": "",    # zero-width non-joiner
        "‍": "",    # zero-width joiner
        "﻿": "",    # byte order mark
        "…": "...", # ellipsis
        "‘": "'",   # left single quote
        "’": "'",   # right single quote
        "“": '"',   # left double quote
        "”": '"',   # right double quote
        "•": "•",   # bullet
    }
    for old, new in replacements.items():
        if old in text:
            text = text.replace(old, new)

    # Unicode superscript characters to HTML <sup>
    sup_map = {
        "⁰": "0", "¹": "1", "²": "2", "³": "3", "⁴": "4",
        "⁵": "5", "⁶": "6", "⁷": "7", "⁸": "8", "⁹": "9",
        "⁺": "+", "⁻": "-", "⁼": "=", "⁽": "(", "⁾": ")",
        "ⁿ": "n", "ⁱ": "i",
    }
    # Unicode subscript characters to HTML <sub>
    sub_map = {
        "₀": "0", "₁": "1", "₂": "2", "₃": "3", "₄": "4",
        "₅": "5", "₆": "6", "₇": "7", "₈": "8", "₉": "9",
        "₊": "+", "₋": "-", "₌": "=", "₍": "(", "₎": ")",
        "ₐ": "a", "ₑ": "e", "ₒ": "o", "ₓ": "x", "ₕ": "h",
        "ₖ": "k", "ₗ": "l", "ₘ": "m", "ₙ": "n", "ₚ": "p",
        "ₛ": "s", "ₜ": "t",
    }
    for char, val in sup_map.items():
        if char in text:
            text = text.replace(char, f"<sup>{val}</sup>")
    for char, val in sub_map.items():
        if char in text:
            text = text.replace(char, f"<sub>{val}</sub>")

    return text


def typeset_math_to_html(raw_math: str) -> str:
    """Typeset LaTeX math and chemical expressions into beautiful HTML formatting.

    Converts formulas like $H_2O$, $CO_2$, $C_6H_{12}O_6$, $x^2+1$ and complex equations
    into proper subscripts, superscripts, arrows, and cleanly spaced operators.
    """
    s = raw_math.strip()

    # 1. Chemical reaction arrows with labels: \xrightarrow[below]{above}
    s = re.sub(r"\\xrightarrow\[([^\]]*)\]\{([^}]*)\}", r" &mdash;[ \2 / \1 ]&rarr; ", s)
    s = re.sub(r"\\xrightarrow\{([^}]*)\}", r" &mdash;[ \1 ]&rarr; ", s)
    s = re.sub(r"\\(?:rightarrow|to|longrightarrow)", " &rarr; ", s)
    s = re.sub(r"\\(?:leftarrow|longleftarrow)", " &larr; ", s)
    s = re.sub(r"\\(?:leftrightarrow|rightleftharpoons)", " &#8644; ", s)

    # 2. Text wrappers: \text{...}, \mathrm{...}
    s = re.sub(r"\\(?:text|mathrm|mathbf|mathit|textsf)\{([^}]*)\}", r" \1 ", s)

    # 3. Common math symbols & operators
    s = re.sub(r"\\times", " &times; ", s)
    s = re.sub(r"\\pm", " &plusmn; ", s)
    s = re.sub(r"\\mp", " &#8723; ", s)
    s = re.sub(r"\\div", " &divide; ", s)
    s = re.sub(r"\\cdot", " &sdot; ", s)
    s = re.sub(r"\\neq", " &ne; ", s)
    s = re.sub(r"\\leq", " &le; ", s)
    s = re.sub(r"\\geq", " &ge; ", s)
    s = re.sub(r"\\approx", " &asymp; ", s)
    s = re.sub(r"\\equiv", " &equiv; ", s)
    s = re.sub(r"\\infty", " &infin; ", s)
    s = re.sub(r"\\Delta", " &Delta; ", s)
    s = re.sub(r"\\pi", " &pi; ", s)
    s = re.sub(r"\\alpha", " &alpha; ", s)
    s = re.sub(r"\\beta", " &beta; ", s)
    s = re.sub(r"\\gamma", " &gamma; ", s)
    s = re.sub(r"\\theta", " &theta; ", s)
    s = re.sub(r"\\lambda", " &lambda; ", s)
    s = re.sub(r"\\mu", " &mu; ", s)
    s = re.sub(r"\\sigma", " &sigma; ", s)
    s = re.sub(r"\\partial", " &part; ", s)
    s = re.sub(r"\\nabla", " &nabla; ", s)
    s = re.sub(r"\\in", " &isin; ", s)
    s = re.sub(r"\\notin", " &notin; ", s)

    # 4. Spacing commands
    s = re.sub(r"\\[,;!]", " ", s)
    s = re.sub(r"\\(?:quad|qquad|enspace|thinspace)", "  ", s)

    # 5. Fractions & roots
    s = re.sub(r"\\frac\{([^}]*)\}\{([^}]*)\}", r"(\1 / \2)", s)
    s = re.sub(r"\\sqrt\{([^}]*)\}", r"&radic;(\1)", s)

    # 6. Subscripts & superscripts (with braces or single alphanumeric character)
    s = re.sub(r"_\{([^}]*)\}", r"<sub>\1</sub>", s)
    s = re.sub(r"_([a-zA-Z0-9])", r"<sub>\1</sub>", s)
    s = re.sub(r"\^\{([^}]*)\}", r"<sup>\1</sup>", s)
    s = re.sub(r"\^([a-zA-Z0-9+–-])", r"<sup>\1</sup>", s)

    # 7. Strip remaining LaTeX macros e.g. \left, \right, brackets
    s = re.sub(r"\\(?:left|right)\b[.|()\[\]]?", "", s)
    s = re.sub(r"\\[a-zA-Z]+", "", s)
    s = re.sub(r"[{}]", "", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s


def parse_inline_spans(text: str) -> List[InlineSpan]:
    """Tokenize markdown inline text into styled spans."""
    if not text:
        return []

    text = normalize_unicode_for_export(text)

    # Combined regex for inline elements:
    # 1. Inline code: `code`
    # 2. Links: [text](url)
    # 3. Math: $latex$ or \(latex\)
    # 4. Bold+Italic: ***text*** or ___text___
    # 5. Bold: **text** or __text__
    # 6. Italic: *text* or _text_
    # 7. Strikethrough: ~~text~~
    pattern = re.compile(
        r"(?P<code>`(?P<code_inner>[^`\n]+)`)|\b"
        r"(?P<link>\[(?P<link_text>[^\]]+)\]\((?P<link_url>[^)]+)\))|"
        r"(?P<math_d>(?<!\w)\$(?!\s)(?!\d+(?:[.,]\d+)?(?:\s|$))(?P<math_inner>[^$\n]+?)(?<!\s)\$(?!\w))|"
        r"(?P<math_p>\\\((?P<math_p_inner>[^)]+)\\\))|"
        r"(?P<bold_italic>(?:\*\*\*|___)(?P<bi_inner>.+?)(?:\*\*\*|___))|"
        r"(?P<bold>(?:\*\*|__)(?P<b_inner>.+?)(?:\*\*|__))|"
        r"(?P<italic>(?:\*|_)(?P<i_inner>.+?)(?:\*|_))|"
        r"(?P<strike>~~(?P<s_inner>.+?)~~)",
        re.DOTALL,
    )

    spans: List[InlineSpan] = []
    last_idx = 0

    for match in pattern.finditer(text):
        start, end = match.span()
        if start > last_idx:
            plain_text = text[last_idx:start]
            if plain_text:
                spans.append(InlineSpan(text=plain_text))

        if match.group("code"):
            spans.append(InlineSpan(text=match.group("code_inner") or "", code=True))
        elif match.group("link"):
            url = (match.group("link_url") or "").strip()
            link_text = match.group("link_text") or url
            spans.append(InlineSpan(text=link_text, link_url=url))
        elif match.group("math_d"):
            spans.append(InlineSpan(text=match.group("math_inner") or "", math=True))
        elif match.group("math_p"):
            spans.append(InlineSpan(text=match.group("math_p_inner") or "", math=True))
        elif match.group("bold_italic"):
            spans.append(InlineSpan(text=match.group("bi_inner") or "", bold=True, italic=True))
        elif match.group("bold"):
            spans.append(InlineSpan(text=match.group("b_inner") or "", bold=True))
        elif match.group("italic"):
            spans.append(InlineSpan(text=match.group("i_inner") or "", italic=True))
        elif match.group("strike"):
            spans.append(InlineSpan(text=match.group("s_inner") or "", strike=True))

        last_idx = end

    if last_idx < len(text):
        tail = text[last_idx:]
        if tail:
            spans.append(InlineSpan(text=tail))

    return spans if spans else [InlineSpan(text=text)]


def inline_spans_to_reportlab_xml(spans: List[InlineSpan]) -> str:
    """Convert inline spans to safe XML markup for ReportLab Paragraphs."""
    out = []
    for span in spans:
        if span.math:
            typeset = typeset_math_to_html(span.text)
            out.append(f'<font color="#b45739"><b>{typeset}</b></font>')
            continue

        escaped = saxutils.escape(span.text)
        # Convert newlines to <br/>
        escaped = escaped.replace("\n", "<br/>")
        if span.code:
            escaped = f'<font name="Courier" color="#141413" backColor="#f0ede6">&nbsp;{escaped}&nbsp;</font>'
        if span.strike:
            escaped = f"<strike>{escaped}</strike>"
        if span.bold:
            escaped = f"<b>{escaped}</b>"
        if span.italic:
            escaped = f"<i>{escaped}</i>"
        if span.link_url:
            safe_url = saxutils.quoteattr(span.link_url)
            escaped = f'<a href={safe_url} color="#d97757"><u>{escaped}</u></a>'
        out.append(escaped)
    return "".join(out)


# --- Block Parser ---

def parse_markdown_to_blocks(markdown_text: str) -> List[Block]:
    """Parse Markdown text into structured blocks."""
    if not markdown_text:
        return []

    lines = markdown_text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    blocks: List[Block] = []
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]
        trimmed = line.strip()

        # 1. Blank line
        if not trimmed:
            i += 1
            continue

        # 2. Horizontal Rule (---, ***, ___)
        if re.match(r"^(?:---+|\*\*\*+|___+)\s*$", trimmed):
            blocks.append(HorizontalRuleBlock())
            i += 1
            continue

        # 3. Fenced Code Block (``` or ~~~)
        fence_match = re.match(r"^(`{3,}|~{3,})\s*(\w*)", trimmed)
        if fence_match:
            fence_char = fence_match.group(1)[0]
            fence_len = len(fence_match.group(1))
            language = fence_match.group(2).lower()
            code_lines = []
            i += 1
            while i < n:
                curr = lines[i]
                close_match = re.match(r"^\s*(`{3,}|~{3,})\s*$", curr)
                if close_match and close_match.group(1)[0] == fence_char and len(close_match.group(1)) >= fence_len:
                    i += 1
                    break
                code_lines.append(curr)
                i += 1
            code_content = "\n".join(code_lines)
            blocks.append(CodeBlock(code=code_content, language=language))
            continue

        # 4. Display Math Block ($$...$$ or \[...\])
        if trimmed.startswith("$$"):
            math_lines = []
            if trimmed.endswith("$$") and len(trimmed) > 2:
                math_lines.append(trimmed[2:-2].strip())
                i += 1
            else:
                math_lines.append(trimmed[2:].strip())
                i += 1
                while i < n:
                    curr = lines[i]
                    if curr.strip().endswith("$$"):
                        math_lines.append(curr.strip()[:-2].strip())
                        i += 1
                        break
                    math_lines.append(curr)
                    i += 1
            blocks.append(MathBlock(latex="\n".join([m for m in math_lines if m]).strip()))
            continue

        if trimmed.startswith("\\["):
            math_lines = []
            if trimmed.endswith("\\]") and len(trimmed) > 2:
                math_lines.append(trimmed[2:-2].strip())
                i += 1
            else:
                math_lines.append(trimmed[2:].strip())
                i += 1
                while i < n:
                    curr = lines[i]
                    if curr.strip().endswith("\\]"):
                        math_lines.append(curr.strip()[:-2].strip())
                        i += 1
                        break
                    math_lines.append(curr)
                    i += 1
            blocks.append(MathBlock(latex="\n".join([m for m in math_lines if m]).strip()))
            continue

        # 5. ATX Headings (# Heading)
        heading_match = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            htext = heading_match.group(2).strip()
            # Strip trailing hashes
            htext = re.sub(r"\s+#+$", "", htext)
            blocks.append(HeadingBlock(level=level, text=htext))
            i += 1
            continue

        # 6. Blockquote (> quote)
        if trimmed.startswith(">"):
            quote_lines = []
            while i < n:
                curr = lines[i].strip()
                if not curr:
                    break
                if curr.startswith(">"):
                    # Remove leading > and optional space
                    quote_lines.append(re.sub(r"^>\s?", "", lines[i]))
                else:
                    # Lazy continuation
                    quote_lines.append(lines[i])
                i += 1
            blocks.append(BlockquoteBlock(lines=quote_lines))
            continue

        # 7. Markdown Tables (| col | col |)
        if "|" in line and i + 1 < n and re.match(r"^\s*\|?[\s:-]+(?:\|[\s:-]+)+\|?\s*$", lines[i + 1]):
            header_line = line
            sep_line = lines[i + 1]
            raw_headers = [c.strip() for c in header_line.strip().strip("|").split("|")]
            raw_seps = [c.strip() for c in sep_line.strip().strip("|").split("|")]
            alignments = []
            for s in raw_seps:
                if s.startswith(":") and s.endswith(":"):
                    alignments.append("center")
                elif s.endswith(":"):
                    alignments.append("right")
                else:
                    alignments.append("left")

            col_count = max(len(raw_headers), len(alignments), 1)
            # Pad headers
            while len(raw_headers) < col_count:
                raw_headers.append("")
            while len(alignments) < col_count:
                alignments.append("left")

            table_rows = []
            i += 2
            while i < n:
                curr = lines[i].strip()
                if not curr or "|" not in curr:
                    break
                row_cells = [c.strip() for c in curr.strip("|").split("|")]
                # Normalize length
                if len(row_cells) < col_count:
                    row_cells.extend([""] * (col_count - len(row_cells)))
                elif len(row_cells) > col_count:
                    row_cells = row_cells[:col_count]
                table_rows.append(row_cells)
                i += 1

            blocks.append(TableBlock(headers=raw_headers, alignments=alignments, rows=table_rows))
            continue

        # 8. Lists (Unordered -/ * / + or Ordered 1. / 2.)
        list_match = re.match(r"^(\s*)([-*+]|\d+\.)\s+(.*)$", line)
        if list_match:
            ordered = list_match.group(2)[0].isdigit()
            items: List[ListItem] = []
            while i < n:
                curr_line = lines[i]
                item_match = re.match(r"^(\s*)([-*+]|\d+\.)\s+(.*)$", curr_line)
                if item_match:
                    indent_spaces = len(item_match.group(1))
                    level = indent_spaces // 2
                    is_num = item_match.group(2)[0].isdigit()
                    num_val = int(item_match.group(2).rstrip(".")) if is_num else None
                    item_text = item_match.group(3)
                    items.append(ListItem(text=item_text, level=min(level, 4), number=num_val))
                    i += 1
                elif items and curr_line.strip() and (curr_line.startswith("  ") or curr_line.startswith("\t")):
                    # Continuation of previous item
                    items[-1].text += " " + curr_line.strip()
                    i += 1
                else:
                    break
            blocks.append(ListBlock(ordered=ordered, items=items))
            continue

        # 9. Regular Paragraph (accumulate consecutive text lines)
        para_lines = [line.strip()]
        i += 1
        while i < n:
            curr = lines[i]
            if not curr.strip():
                break
            if (
                curr.strip().startswith("#")
                or curr.strip().startswith(">")
                or curr.strip().startswith("```")
                or curr.strip().startswith("~~~")
                or curr.strip().startswith("$$")
                or curr.strip().startswith("\\[")
                or re.match(r"^(?:---+|\*\*\*+|___+)\s*$", curr.strip())
                or re.match(r"^(\s*)([-*+]|\d+\.)\s+", curr)
                or ("|" in curr and i + 1 < n and re.match(r"^\s*\|?[\s:-]+(?:\|[\s:-]+)+\|?\s*$", lines[i + 1]))
            ):
                break
            para_lines.append(curr.strip())
            i += 1

        full_para = " ".join(para_lines)
        if full_para:
            blocks.append(ParagraphBlock(text=full_para))

    return blocks


# --- Filename Sanitization ---

def sanitize_export_filename(title: Optional[str], ext: str) -> str:
    """Generate a clean, safe, URL-friendly filename with extension."""
    ext = ext.lstrip(".").lower()
    if not ext or ext not in {"md", "pdf", "docx"}:
        ext = "md"

    raw_title = (title or "").strip()
    if not raw_title:
        raw_title = "Bimo AI response"

    # Normalize unicode to ASCII
    norm = unicodedata.normalize("NFKD", raw_title)
    ascii_title = norm.encode("ascii", "ignore").decode("ascii")

    # Replace path separators, dots, unsafe chars with hyphens
    slug = re.sub(r"[^\w\s-]", "", ascii_title).strip()
    slug = re.sub(r"[\s_]+", "-", slug).lower()
    slug = re.sub(r"-+", "-", slug).strip("-")

    if not slug:
        slug = "bimo-ai-response"

    # Clamp length
    slug = slug[:80].rstrip("-")
    return f"{slug}.{ext}"


# --- Canonical Markdown Exporter ---

def export_canonical_markdown(title: Optional[str], markdown_content: str, generated_at: Optional[datetime] = None) -> bytes:
    """Return UTF-8 bytes for the canonical Markdown document."""
    dt = generated_at or datetime.now(timezone.utc)
    date_str = dt.strftime("%B %d, %Y at %I:%M %p UTC")
    clean_title = (title or "").strip() or "Bimo AI response"

    # If content already starts with the exact H1 title, don't duplicate
    has_matching_h1 = markdown_content.strip().startswith(f"# {clean_title}")
    header_parts = []
    if not has_matching_h1:
        header_parts.append(f"# {clean_title}\n")
    header_parts.append(f"*Generated on {date_str} · Created with Bimo*\n")
    header_parts.append("---\n")

    full_md = "\n".join(header_parts) + "\n" + markdown_content.strip() + "\n"
    return full_md.encode("utf-8")


# --- ReportLab Text-Based PDF Exporter ---

def export_pdf(title: Optional[str], markdown_content: str, generated_at: Optional[datetime] = None) -> bytes:
    """Generate a high-quality, text-based PDF document using ReportLab."""
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.pdfgen import canvas
        from reportlab.platypus import (
            HRFlowable,
            KeepTogether,
            Paragraph,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        )
    except ImportError as exc:
        raise RuntimeError("reportlab is required for PDF export but not installed") from exc

    dt = generated_at or datetime.now(timezone.utc)
    date_str = dt.strftime("%B %d, %Y at %I:%M %p UTC")
    clean_title = (title or "").strip() or "Bimo AI response"

    # Numbered Canvas for running Header and Footer ("Page X of Y")
    class BimoNumberedCanvas(canvas.Canvas):
        def __init__(self, *args, **kwargs):
            super().__init__(*args, **kwargs)
            self._saved_page_states = []

        def showPage(self):
            self._saved_page_states.append(dict(self.__dict__))
            self._startPage()

        def save(self):
            num_pages = len(self._saved_page_states)
            for state in self._saved_page_states:
                self.__dict__.update(state)
                self.draw_decorations(num_pages)
                super().showPage()
            super().save()

        def draw_decorations(self, page_count):
            self.saveState()
            page_w, page_h = letter
            margin = 48

            # Header (Pages 2+)
            if self._pageNumber > 1:
                self.setFont("Helvetica-Bold", 8)
                self.setFillColor(colors.HexColor("#d97757"))
                self.drawString(margin, page_h - 32, "Bimo")
                self.setFont("Helvetica", 8)
                self.setFillColor(colors.HexColor("#6c6a64"))
                header_title = clean_title if len(clean_title) < 45 else clean_title[:42] + "…"
                self.drawString(margin + 26, page_h - 32, f"· {header_title}")
                self.setStrokeColor(colors.HexColor("#e6dfd8"))
                self.setLineWidth(0.5)
                self.line(margin, page_h - 38, page_w - margin, page_h - 38)

            # Footer (All pages)
            self.setStrokeColor(colors.HexColor("#e6dfd8"))
            self.setLineWidth(0.5)
            self.line(margin, 42, page_w - margin, 42)

            self.setFont("Helvetica", 8)
            self.setFillColor(colors.HexColor("#8e8b82"))
            self.drawString(margin, 30, "AI-generated document · Markdown source · Created with Bimo")

            page_str = f"Page {self._pageNumber} of {page_count}"
            self.drawRightString(page_w - margin, 30, page_str)
            self.restoreState()

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=letter,
        leftMargin=48,
        rightMargin=48,
        topMargin=48,
        bottomMargin=48,
        title=clean_title,
        author="Bimo AI",
    )

    getSampleStyleSheet()

    # Define color palette matching Bimo's design system
    c_primary = colors.HexColor("#d97757")
    c_ink = colors.HexColor("#141413")
    c_body = colors.HexColor("#3d3d3a")
    c_muted = colors.HexColor("#6c6a64")
    c_surface_soft = colors.HexColor("#f5f0e8")
    c_hairline = colors.HexColor("#e6dfd8")
    c_dark_header = colors.HexColor("#2c2a26")

    # Typography Styles
    title_style = ParagraphStyle(
        "BimoTitle",
        fontName="Helvetica-Bold",
        fontSize=20,
        leading=24,
        textColor=c_ink,
        spaceAfter=4,
        keepWithNext=True,
    )
    meta_style = ParagraphStyle(
        "BimoMeta",
        fontName="Helvetica-Oblique",
        fontSize=9,
        leading=12,
        textColor=c_muted,
        spaceAfter=12,
        keepWithNext=True,
    )
    h1_style = ParagraphStyle(
        "BimoH1",
        fontName="Helvetica-Bold",
        fontSize=15,
        leading=19,
        textColor=c_ink,
        spaceBefore=14,
        spaceAfter=6,
        keepWithNext=True,
    )
    h2_style = ParagraphStyle(
        "BimoH2",
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=17,
        textColor=c_ink,
        spaceBefore=12,
        spaceAfter=5,
        keepWithNext=True,
    )
    h3_style = ParagraphStyle(
        "BimoH3",
        fontName="Helvetica-Bold",
        fontSize=11.5,
        leading=15,
        textColor=c_primary,
        spaceBefore=10,
        spaceAfter=4,
        keepWithNext=True,
    )
    body_style = ParagraphStyle(
        "BimoBody",
        fontName="Helvetica",
        fontSize=9.5,
        leading=14,
        textColor=c_body,
        spaceAfter=8,
    )
    list_item_style = ParagraphStyle(
        "BimoListItem",
        fontName="Helvetica",
        fontSize=9.5,
        leading=13.5,
        textColor=c_body,
        spaceAfter=4,
    )
    code_text_style = ParagraphStyle(
        "BimoCodeText",
        fontName="Courier",
        fontSize=8.5,
        leading=11.5,
        textColor=c_ink,
    )
    quote_text_style = ParagraphStyle(
        "BimoQuoteText",
        fontName="Helvetica-Oblique",
        fontSize=9.5,
        leading=14,
        textColor=c_body,
    )
    table_cell_style = ParagraphStyle(
        "BimoTableCell",
        fontName="Helvetica",
        fontSize=8.5,
        leading=11.5,
        textColor=c_body,
    )
    table_head_style = ParagraphStyle(
        "BimoTableHead",
        fontName="Helvetica-Bold",
        fontSize=8.5,
        leading=11.5,
        textColor=colors.white,
    )

    story = []

    # Title & Metadata
    clean_title_norm = normalize_unicode_for_export(clean_title)
    # If the markdown content already starts with an H1 (# Title), avoid duplicating it as the body heading
    blocks = parse_markdown_to_blocks(markdown_content)
    if blocks and isinstance(blocks[0], HeadingBlock) and blocks[0].level == 1:
        first_h1_norm = normalize_unicode_for_export(blocks[0].text).strip()
        if first_h1_norm.lower() == clean_title_norm.lower() or not clean_title_norm:
            clean_title_norm = first_h1_norm
            blocks.pop(0)

    story.append(Paragraph(saxutils.escape(clean_title_norm), title_style))
    story.append(Paragraph(saxutils.escape(f"Generated on {date_str} · Bimo AI Assistant"), meta_style))
    story.append(HRFlowable(width="100%", thickness=0.75, color=c_hairline, spaceAfter=14))

    content_width = letter[0] - 96  # 612 - 96 = 516 pt

    for block in blocks:
        if isinstance(block, HeadingBlock):
            spans = parse_inline_spans(block.text)
            xml_text = inline_spans_to_reportlab_xml(spans)
            if block.level == 1:
                story.append(Paragraph(xml_text, h1_style))
            elif block.level == 2:
                story.append(Paragraph(xml_text, h2_style))
            else:
                story.append(Paragraph(xml_text, h3_style))

        elif isinstance(block, ParagraphBlock):
            spans = parse_inline_spans(block.text)
            xml_text = inline_spans_to_reportlab_xml(spans)
            story.append(Paragraph(xml_text, body_style))

        elif isinstance(block, HorizontalRuleBlock):
            story.append(Spacer(1, 4))
            story.append(HRFlowable(width="100%", thickness=0.5, color=c_hairline, spaceAfter=8, spaceBefore=4))

        elif isinstance(block, CodeBlock):
            # Render code block in a shaded container with header
            lang_label = block.language.upper() if block.language else "CODE"
            escaped_code = saxutils.escape(block.code).replace("\n", "<br/>").replace(" ", "&nbsp;")
            code_para = Paragraph(escaped_code, code_text_style)

            head_style = ParagraphStyle(
                "CodeHead",
                fontName="Helvetica-Bold",
                fontSize=7.5,
                leading=9,
                textColor=c_muted,
            )
            head_para = Paragraph(saxutils.escape(lang_label), head_style)

            tbl_data = [[head_para], [code_para]]
            code_tbl = Table(tbl_data, colWidths=[content_width])
            code_tbl.setStyle(
                TableStyle([
                    ("BACKGROUND", (0, 0), (-1, -1), c_surface_soft),
                    ("BOX", (0, 0), (-1, -1), 0.75, c_hairline),
                    ("LINEBELOW", (0, 0), (-1, 0), 0.5, c_hairline),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                    ("LEFTPADDING", (0, 0), (-1, -1), 8),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ])
            )
            story.append(Spacer(1, 3))
            story.append(KeepTogether([code_tbl]))
            story.append(Spacer(1, 8))

        elif isinstance(block, MathBlock):
            math_html = typeset_math_to_html(block.latex)
            math_p = Paragraph(f'<font color="#b45739"><b>{math_html}</b></font>', ParagraphStyle(
                "MathBlockStyle",
                fontName="Helvetica",
                fontSize=11,
                leading=15,
                alignment=1,  # Centered
            ))
            math_tbl = Table([[math_p]], colWidths=[content_width])
            math_tbl.setStyle(
                TableStyle([
                    ("BACKGROUND", (0, 0), (-1, -1), c_surface_soft),
                    ("BOX", (0, 0), (-1, -1), 0.75, c_hairline),
                    ("TOPPADDING", (0, 0), (-1, -1), 8),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                    ("LEFTPADDING", (0, 0), (-1, -1), 10),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ])
            )
            story.append(Spacer(1, 3))
            story.append(KeepTogether([math_tbl]))
            story.append(Spacer(1, 8))

        elif isinstance(block, BlockquoteBlock):
            quote_content = "<br/>".join([saxutils.escape(line) for line in block.lines])
            quote_para = Paragraph(quote_content, quote_text_style)
            quote_tbl = Table([[quote_para]], colWidths=[content_width])
            quote_tbl.setStyle(
                TableStyle([
                    ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#faf8f5")),
                    ("LINEBEFORE", (0, 0), (0, -1), 3.0, c_primary),
                    ("TOPPADDING", (0, 0), (-1, -1), 6),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                    ("LEFTPADDING", (0, 0), (-1, -1), 10),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ])
            )
            story.append(Spacer(1, 3))
            story.append(quote_tbl)
            story.append(Spacer(1, 8))

        elif isinstance(block, ListBlock):
            for idx, item in enumerate(block.items):
                indent = 12 + (item.level * 14)
                spans = parse_inline_spans(item.text)
                xml_text = inline_spans_to_reportlab_xml(spans)

                if item.number is not None:
                    bullet_prefix = f"{item.number}."
                elif block.ordered:
                    bullet_prefix = f"{idx + 1}."
                else:
                    bullet_prefix = "•"

                p_style = ParagraphStyle(
                    f"ListP_{item.level}",
                    parent=list_item_style,
                    leftIndent=indent,
                    firstLineIndent=-12,
                )
                story.append(Paragraph(f'<font color="#d97757">{bullet_prefix}</font>&nbsp;&nbsp;{xml_text}', p_style))
            story.append(Spacer(1, 6))

        elif isinstance(block, TableBlock):
            col_count = len(block.headers) if block.headers else (len(block.rows[0]) if block.rows else 1)
            if col_count == 0:
                continue

            col_width = content_width / col_count
            tbl_rows_data = []

            # Headers
            if block.headers:
                head_row = []
                for h in block.headers:
                    h_spans = parse_inline_spans(h)
                    h_xml = inline_spans_to_reportlab_xml(h_spans)
                    head_row.append(Paragraph(h_xml, table_head_style))
                tbl_rows_data.append(head_row)

            # Rows
            for r in block.rows:
                cell_row = []
                for c in r:
                    c_spans = parse_inline_spans(c)
                    c_xml = inline_spans_to_reportlab_xml(c_spans)
                    cell_row.append(Paragraph(c_xml, table_cell_style))
                tbl_rows_data.append(cell_row)

            if tbl_rows_data:
                table = Table(tbl_rows_data, colWidths=[col_width] * col_count, repeatRows=1 if block.headers else 0)
                t_style = [
                    ("BOX", (0, 0), (-1, -1), 0.75, c_hairline),
                    ("INNERGRID", (0, 0), (-1, -1), 0.5, c_hairline),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ]
                if block.headers:
                    t_style.append(("BACKGROUND", (0, 0), (-1, 0), c_dark_header))
                # Alternating zebra background
                start_row = 1 if block.headers else 0
                for r_idx in range(start_row, len(tbl_rows_data)):
                    if (r_idx - start_row) % 2 == 1:
                        t_style.append(("BACKGROUND", (0, r_idx), (-1, r_idx), c_surface_soft))

                table.setStyle(TableStyle(t_style))
                story.append(Spacer(1, 4))
                story.append(table)
                story.append(Spacer(1, 8))

    doc.build(story, canvasmaker=BimoNumberedCanvas)
    return buf.getvalue()


# --- python-docx Word Document Exporter ---

def export_docx(title: Optional[str], markdown_content: str, generated_at: Optional[datetime] = None) -> bytes:
    """Generate an editable Microsoft Word document (.docx) using python-docx."""
    try:
        import docx
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:
        raise RuntimeError("python-docx is required for DOCX export but not installed") from exc

    dt = generated_at or datetime.now(timezone.utc)
    date_str = dt.strftime("%B %d, %Y at %I:%M %p UTC")
    clean_title = (title or "").strip() or "Bimo AI response"

    doc = docx.Document()

    # 1-inch margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

        # Header
        header = section.header
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run(f"Bimo · {clean_title}")
        hrun.font.name = "Arial"
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGBColor(142, 139, 130)

        # Footer
        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        frun = fp.add_run("AI-generated document · Markdown source · Created with Bimo")
        frun.font.name = "Arial"
        frun.font.size = Pt(8.5)
        frun.font.color.rgb = RGBColor(142, 139, 130)

    # Base styling
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(61, 61, 58)

    # Helper: add styled runs from inline spans
    def add_spans_to_paragraph(p, spans: List[InlineSpan]):
        for span in spans:
            if span.link_url:
                add_hyperlink(p, span.link_url, span.text)
                continue

            run = p.add_run(span.text)
            if span.bold:
                run.bold = True
            if span.italic:
                run.italic = True
            if span.strike:
                run.font.strike = True
            if span.code:
                run.font.name = "Consolas"
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(20, 20, 19)
                # Add background shading element to run
                rPr = run._r.get_or_add_rPr()
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0EDE6"/>')
                rPr.append(shd)
            elif span.math:
                run.font.name = "Consolas"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(180, 87, 57)
                run.bold = True

    def add_hyperlink(paragraph, url: str, text: str):
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        hyperlink = parse_xml(
            f'<w:hyperlink {nsdecls("w")} r:id="{r_id}" w:history="1">'
            f'<w:r><w:rPr><w:color w:val="D97757"/><w:u w:val="single"/></w:rPr>'
            f'<w:t>{saxutils.escape(text)}</w:t></w:r></w:hyperlink>'
        )
        paragraph._p.append(hyperlink)

    # Document Title & Metadata
    clean_title_norm = normalize_unicode_for_export(clean_title)
    blocks = parse_markdown_to_blocks(markdown_content)
    if blocks and isinstance(blocks[0], HeadingBlock) and blocks[0].level == 1:
        first_h1_norm = normalize_unicode_for_export(blocks[0].text).strip()
        if first_h1_norm.lower() == clean_title_norm.lower() or not clean_title_norm:
            clean_title_norm = first_h1_norm
            blocks.pop(0)

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(2)
    title_run = title_p.add_run(clean_title_norm)
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(20, 20, 19)

    meta_p = doc.add_paragraph()
    meta_p.paragraph_format.space_before = Pt(0)
    meta_p.paragraph_format.space_after = Pt(12)
    meta_run = meta_p.add_run(f"Generated on {date_str} · Bimo AI Assistant")
    meta_run.italic = True
    meta_run.font.size = Pt(9.5)
    meta_run.font.color.rgb = RGBColor(108, 106, 100)

    # Divider line
    div_p = doc.add_paragraph()
    div_p.paragraph_format.space_before = Pt(0)
    div_p.paragraph_format.space_after = Pt(14)
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="E6DFD8"/></w:pBdr>')
    div_p._p.get_or_add_pPr().append(pBdr)

    for block in blocks:
        if isinstance(block, HeadingBlock):
            p = doc.add_paragraph()
            p.paragraph_format.keep_with_next = True
            if block.level == 1:
                p.paragraph_format.space_before = Pt(16)
                p.paragraph_format.space_after = Pt(6)
                spans = parse_inline_spans(block.text)
                for span in spans:
                    r = p.add_run(span.text)
                    r.bold = True
                    r.font.size = Pt(15)
                    r.font.color.rgb = RGBColor(20, 20, 19)
            elif block.level == 2:
                p.paragraph_format.space_before = Pt(13)
                p.paragraph_format.space_after = Pt(5)
                spans = parse_inline_spans(block.text)
                for span in spans:
                    r = p.add_run(span.text)
                    r.bold = True
                    r.font.size = Pt(13)
                    r.font.color.rgb = RGBColor(20, 20, 19)
            else:
                p.paragraph_format.space_before = Pt(11)
                p.paragraph_format.space_after = Pt(4)
                spans = parse_inline_spans(block.text)
                for span in spans:
                    r = p.add_run(span.text)
                    r.bold = True
                    r.font.size = Pt(11.5)
                    r.font.color.rgb = RGBColor(217, 119, 87)

        elif isinstance(block, ParagraphBlock):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(7)
            p.paragraph_format.line_spacing = 1.2
            spans = parse_inline_spans(block.text)
            add_spans_to_paragraph(p, spans)

        elif isinstance(block, HorizontalRuleBlock):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(10)
            pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="E6DFD8"/></w:pBdr>')
            p._p.get_or_add_pPr().append(pBdr)

        elif isinstance(block, CodeBlock):
            # Single cell shaded table for code block
            tbl = doc.add_table(rows=1, cols=1)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            tbl.autofit = False
            cell = tbl.cell(0, 0)
            cell.width = Inches(6.5)

            # Shading & borders
            tcPr = cell._tc.get_or_add_tcPr()
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F0E8"/>')
            tcBorders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'<w:top w:val="single" w:sz="4" w:color="E6DFD8"/>'
                f'<w:left w:val="single" w:sz="4" w:color="E6DFD8"/>'
                f'<w:bottom w:val="single" w:sz="4" w:color="E6DFD8"/>'
                f'<w:right w:val="single" w:sz="4" w:color="E6DFD8"/>'
                f'</w:tcBorders>'
            )
            tcMar = parse_xml(
                f'<w:tcMar {nsdecls("w")}>'
                f'<w:top w:w="120" w:type="dxa"/>'
                f'<w:bottom w:w="120" w:type="dxa"/>'
                f'<w:left w:w="160" w:type="dxa"/>'
                f'<w:right w:w="160" w:type="dxa"/>'
                f'</w:tcMar>'
            )
            tcPr.append(shd)
            tcPr.append(tcBorders)
            tcPr.append(tcMar)

            cp = cell.paragraphs[0]
            cp.paragraph_format.space_before = Pt(0)
            cp.paragraph_format.space_after = Pt(0)
            if block.language:
                hrun = cp.add_run(f"// {block.language.upper()}\n")
                hrun.font.name = "Consolas"
                hrun.font.size = Pt(8.5)
                hrun.font.color.rgb = RGBColor(108, 106, 100)
                hrun.bold = True

            crun = cp.add_run(block.code)
            crun.font.name = "Consolas"
            crun.font.size = Pt(9.0)
            crun.font.color.rgb = RGBColor(20, 20, 19)

            # Spacing after table
            post_p = doc.add_paragraph()
            post_p.paragraph_format.space_before = Pt(0)
            post_p.paragraph_format.space_after = Pt(6)

        elif isinstance(block, MathBlock):
            tbl = doc.add_table(rows=1, cols=1)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = tbl.cell(0, 0)
            cell.width = Inches(6.5)

            tcPr = cell._tc.get_or_add_tcPr()
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F0E8"/>')
            tcBorders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'<w:top w:val="single" w:sz="4" w:color="E6DFD8"/>'
                f'<w:left w:val="single" w:sz="4" w:color="E6DFD8"/>'
                f'<w:bottom w:val="single" w:sz="4" w:color="E6DFD8"/>'
                f'<w:right w:val="single" w:sz="4" w:color="E6DFD8"/>'
                f'</w:tcBorders>'
            )
            tcPr.append(shd)
            tcPr.append(tcBorders)

            mp = cell.paragraphs[0]
            mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            mrun = mp.add_run(block.latex)
            mrun.font.name = "Consolas"
            mrun.font.size = Pt(10)
            mrun.bold = True
            mrun.font.color.rgb = RGBColor(180, 87, 57)

            post_p = doc.add_paragraph()
            post_p.paragraph_format.space_before = Pt(0)
            post_p.paragraph_format.space_after = Pt(6)

        elif isinstance(block, BlockquoteBlock):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(6)
            pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="18" w:space="8" w:color="D97757"/></w:pBdr>')
            p._p.get_or_add_pPr().append(pBdr)

            full_quote = "\n".join(block.lines)
            spans = parse_inline_spans(full_quote)
            add_spans_to_paragraph(p, spans)

        elif isinstance(block, ListBlock):
            for idx, item in enumerate(block.items):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25 + (item.level * 0.2))
                p.paragraph_format.first_line_indent = Inches(-0.15)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(3)

                if item.number is not None:
                    prefix = f"{item.number}. "
                elif block.ordered:
                    prefix = f"{idx + 1}. "
                else:
                    prefix = "• "

                brun = p.add_run(prefix)
                brun.bold = True
                brun.font.color.rgb = RGBColor(217, 119, 87)

                spans = parse_inline_spans(item.text)
                add_spans_to_paragraph(p, spans)

        elif isinstance(block, TableBlock):
            col_count = len(block.headers) if block.headers else (len(block.rows[0]) if block.rows else 1)
            row_count = (1 if block.headers else 0) + len(block.rows)
            if row_count == 0 or col_count == 0:
                continue

            tbl = doc.add_table(rows=row_count, cols=col_count)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

            curr_r = 0
            if block.headers:
                for c_idx, h_text in enumerate(block.headers[:col_count]):
                    cell = tbl.cell(0, c_idx)
                    # Header cell shading: Dark Slate #2C2A26
                    tcPr = cell._tc.get_or_add_tcPr()
                    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="2C2A26"/>')
                    tcPr.append(shd)
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after = Pt(4)
                    spans = parse_inline_spans(h_text)
                    for span in spans:
                        r = p.add_run(span.text)
                        r.bold = True
                        r.font.size = Pt(9.5)
                        r.font.color.rgb = RGBColor(255, 255, 255)
                curr_r = 1

            for r_idx, row_data in enumerate(block.rows):
                for c_idx, cell_text in enumerate(row_data[:col_count]):
                    cell = tbl.cell(curr_r + r_idx, c_idx)
                    if r_idx % 2 == 1:
                        tcPr = cell._tc.get_or_add_tcPr()
                        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F9F8F5"/>')
                        tcPr.append(shd)
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(3)
                    p.paragraph_format.space_after = Pt(3)
                    spans = parse_inline_spans(cell_text)
                    add_spans_to_paragraph(p, spans)

            # Set light grid borders on table
            tblBorders = parse_xml(
                f'<w:tblBorders {nsdecls("w")}>'
                f'<w:top w:val="single" w:sz="4" w:color="E6DFD8"/>'
                f'<w:left w:val="single" w:sz="4" w:color="E6DFD8"/>'
                f'<w:bottom w:val="single" w:sz="4" w:color="E6DFD8"/>'
                f'<w:right w:val="single" w:sz="4" w:color="E6DFD8"/>'
                f'<w:insideH w:val="single" w:sz="4" w:color="E6DFD8"/>'
                f'<w:insideV w:val="single" w:sz="4" w:color="E6DFD8"/>'
                f'</w:tblBorders>'
            )
            tbl._tbl.tblPr.append(tblBorders)

            post_p = doc.add_paragraph()
            post_p.paragraph_format.space_before = Pt(0)
            post_p.paragraph_format.space_after = Pt(6)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
